import re
import pandas as pd
from docx import Document
from docx.shared import Inches
import os
import urllib.parse
import argparse
import json

def read_md_file(file_path):

    try:

        # 使用UTF-8编码打开文件
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            print("[+]读取{"+file_path+"}成功")
            #print(content)
            return content
    except Exception as e:
        print(f"读取文件失败: {str(e)}")
        return None


def parse_md_string(md_string):
    # 初始化结果字典
    result = {
        '任务信息': {
            '公司名称': '',
            '系统名称': '',
            '域名': '',
            '漏洞名称': '',
            '漏洞URL': '',
            '危害级别': '',
            '日期': '',
            '测试人员': ''
        },
        '漏洞详情': {
            '数据包': '',
            '复现过程': ''
        }
    }

    # 分割Markdown内容为不同的部分
    sections = re.split(r'(?m)^###\s+', md_string)
    sections = [s.strip() for s in sections if s.strip()]

    # 处理每个部分
    for section in sections:
        lines = section.split('\n')
        title = lines[0].strip()
        content = '\n'.join(lines[1:]).strip()

        if title == '任务信息':
            print("[+]开始读取任务信息")
            # 提取表格内容
            table_pattern = r'\|\s*(.*?)\s*\|\s*(.*?)\s*\|'
            matches = re.findall(table_pattern, content)
            for key, value in matches:
                if key in ['公司名称', '系统名称', '域名', '漏洞名称', '漏洞URL', '危害级别', '日期', '测试人员']:
                    if value == '':
                        print("[-]"+key+'缺失')
                    #print(key+":"+value)
                    result['任务信息'][key] = value



        elif title == result.get('任务信息', {}).get('漏洞名称'):
            print("[+]正在读取漏洞复现过程")
            # 提取URL内容（标题后的第一部分文本）
            code_block = ''
            body_content = ''

            # 分割内容为代码块前后两部分
            code_block_pattern = r'```([\s\S]*?)```'
            code_matches = re.search(code_block_pattern, content)

            if code_matches:
                code_start = content.find('```')
                code_end = content.find('```', code_start + 3) + 3

                # 代码块内容
                code_block = code_matches.group(1).strip()

                #print(code_block)
                # 剩余内容
                body_content = content[code_end:].strip()
                #print(body_content)
                # 匹配并移除![...](...)格式的图片标记
                body_content = re.sub(r'!\[.*?\]\((.*?)\)', r'\1', body_content)
                # 将连续的换行符替换为单个换行符
                body_content = re.sub(r'\n{2,}','\n',body_content)
                # 去掉分割符
                body_content = body_content.replace("---", "")


            result['漏洞详情']['数据包'] = code_block
            result['漏洞详情']['复现过程'] = body_content
            if result['漏洞详情']['数据包'] != '' and result['漏洞详情']['复现过程']:
                print("[+]读取到数据包和复现过程")

    #print(result)
    return result


def find_vulnerability_in_excel(result, excel_file_path):
    """
    在Excel文件中查找漏洞名称并返回对应行的内容

    参数:
    result (dict): 包含任务信息的字典，通过parse_md_string函数生成
    excel_file_path (str): Excel文件的路径

    返回:
    pandas.DataFrame: 包含匹配行的DataFrame，如果未找到则返回空DataFrame
    """
    # 读取Excel文件
    try:
        df = pd.read_excel(excel_file_path)
        print("[+]读取{"+excel_file_path+"}成功")
    except Exception as e:
        print(f"读取Excel文件时出错: {e}")
        return pd.DataFrame()

    # 获取漏洞名称
    vulnerability_name = result.get('任务信息', {}).get('漏洞名称')
    if not vulnerability_name:
        print("未找到漏洞名称")
        return pd.DataFrame()

    # 在Excel文件中查找漏洞名称
    if '漏洞名称' in df.columns:
        # 查找完全匹配的行
        matching_rows = df[df['漏洞名称'] == vulnerability_name]

        if not matching_rows.empty:
            print(f"[+]找到 {len(matching_rows)} 条匹配记录")
            print("[+]匹配到 *"+vulnerability_name+"* 漏洞")
            # 提取威胁描述和解决方案字段
            threat_description = matching_rows['威胁描述'].iloc[0] if '威胁描述' in matching_rows.columns else ''
            solution = matching_rows['解决方案'].iloc[0] if '解决方案' in matching_rows.columns else ''
            # print(threat_description)
            # print(solution)
            # 添加到result的漏洞名部分
            if '漏洞详情' in result:
                result['漏洞详情']['威胁描述'] = threat_description
                result['漏洞详情']['解决方案'] = solution
                print("[+]已将威胁描述和解决方案添加到结果中")

            return matching_rows
        else:
            print(f"未找到漏洞名称为 '{vulnerability_name}' 的记录")


    return pd.DataFrame()

def flatten_dict(nested_dict, prefix=''):
    """
    将嵌套字典转换为平面替换字典（无前缀）
    :param nested_dict: 嵌套字典
    :param prefix: 内部使用参数，无需传递
    :return: 平面字典 {'{{键名}}': 值}
    """
    flat_dict = {}
    for key, value in nested_dict.items():
        if isinstance(value, dict):
            flat_dict.update(flatten_dict(value))
        else:
            # # 处理Windows路径中的反斜杠
            # if isinstance(value, str) and '\\' in value:
            #     value = value.replace('\\', '/')
            flat_dict[f"{{{key}}}"] = value  # 改为单层花括号更符合常见模板格式
    return flat_dict


def replace_in_word_template(template_path, output_path, replacements):
    """
    完整的Word模板处理函数，包含：
    1. URL编码解码（如%20转空格）
    2. 智能占位符检测
    3. 图片插入功能
    4. 空内容跳过优化
    """
    try:
	
        output_dir = os.path.dirname(output_path)
        if output_dir:  # 非空目录时才创建
            os.makedirs(output_dir, exist_ok=True)
        doc = Document(template_path)

        # ================= 预处理复现过程内容 =================
        def extract_text_paths(content):
            """从纯文本中提取图片路径"""
            # 匹配包含扩展名的路径（支持URL编码）
            pattern = r'''
                        (?:^|\n)                    # 行首或换行后
                        (                           # 捕获组开始
                            (?:[a-zA-Z]:)?          # Windows盘符（可选）
                            (?:                     # 路径部分
                                [\\/]              # 路径分隔符
                                (?:               
                                    [^\\/\s]        # 非分隔符字符
                                    |%[0-9a-fA-F]{2}  # URL编码字符
                                )+
                            )+                    
                            \.(?:png|jpg|jpeg|gif)  # 图片扩展名
                        )                           # 捕获组结束
                        (?=\s|$|\n)                 # 后跟空白或行尾
                    '''
            matches = re.findall(pattern, content, re.VERBOSE)

            # 解码并标准化路径
            processed_paths = []
            for raw_path in matches:
                try:
                    # 解码URL编码（如%20转空格）
                    decoded_path = urllib.parse.unquote(raw_path)
                    # 标准化路径格式
                    normalized_path = os.path.normpath(decoded_path)
                    processed_paths.append(normalized_path)
                    print("[+]读取图片成功 "+normalized_path)
                except Exception as e:
                    print(f"[-] 路径处理失败: {raw_path} ({str(e)})")

            # 生成替换后的文本（保留原始结构）
            replace_text = re.sub(
                pattern,
                '\n[图片位置]',  # 用占位符标记图片位置
                content,
                flags=re.VERBOSE
            )

            return processed_paths, replace_text

        # 处理复现过程内容
        reproduce_content = replacements.get('{复现过程}', '')
        image_paths, processed_text = extract_text_paths(reproduce_content)

        # ================= 段落处理核心逻辑 =================
        def process_paragraph(paragraph):
            """增强型段落处理函数"""
            original_text = paragraph.text

            # 跳过空段落
            if not original_text.strip():
                return

            # 特殊处理复现过程段落
            if '{复现过程}' in original_text:
                # 清空原有内容
                for run in paragraph.runs:
                    run.text = ''

                # 分割处理后的文本为多行
                lines = processed_text.split('\n')
                img_index = 0  # 图片索引

                current_run = paragraph.add_run()  # 初始化当前Run
                for line in lines:
                    # 处理包含占位符的行
                    if '[图片位置]' in line:
                        # 分割占位符前后的文本
                        parts = line.split('[图片位置]', 1)  # 最多分割一次

                        # 添加前半部分文本
                        if parts[0].strip():
                            current_run.add_text(parts[0])

                        # 插入图片
                        if img_index < len(image_paths):
                            img_path = image_paths[img_index]
                            try:
                                if os.path.exists(img_path):
                                    # 新建Run插入图片（保持格式独立）
                                    pic_run = paragraph.add_run()
                                    pic_run.add_picture(img_path, width=Inches(5.5))
                                    img_index += 1
                                    # 新建Run继续添加文本
                                    current_run = paragraph.add_run()
                            except Exception as e:
                                error_run = paragraph.add_run()
                                error_run.add_text(f" [图片加载失败: {str(e)}] ")
                                current_run = error_run

                        # 添加后半部分文本
                        if len(parts) > 1 and parts[1].strip():
                            current_run.add_text(parts[1])

                        # 添加换行符
                        current_run.add_text('\n')
                        current_run = paragraph.add_run()

                    # 处理普通文本行
                    else:
                        if line.strip():
                            current_run.add_text(line + '\n')
                            current_run = paragraph.add_run()
                return

            # 普通占位符处理（优化版）
            if '{' not in original_text and '}' not in original_text:
                return

            # ...（原有普通占位符处理逻辑保持不变）...
            # 精确提取占位符
            placeholders = re.findall(r'\{[^{}]*\}', original_text)
            if not placeholders:
                return

            new_text = original_text
            replaced = False

            # 仅替换存在的占位符
            for ph in set(placeholders):
                if ph in replacements:
                    replacement = urllib.parse.unquote(str(replacements[ph]))  # 解码替换值
                    new_text = new_text.replace(ph, replacement)
                    replaced = True

            if replaced:
                # 保留第一个Run的格式
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    # 清空后续Runs
                    for run in paragraph.runs[1:]:
                        run.text = ''
                else:
                    paragraph.add_run(new_text)
        # ================= 文档遍历处理 =================
        # 处理所有段落
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph)

        # 处理表格（逻辑保持不变）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if not cell.text.strip():
                        continue
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph)

        doc.save(output_path)

        print(f"[+] 文档已生成：{output_path}")
        return True

    except Exception as e:
        print(f"[-] 处理失败：{str(e)}")
        return False

def read_config(config_path='config.json'):
    """
    读取配置文件，获取 xls_file_path 和 word_template_path
    """
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)
            # 验证必要字段存在
            if 'xls_file_path' not in config or 'word_template_path' not in config:
                raise ValueError("配置文件中缺少必要字段 (xls_file_path 或 word_template_path)")
            return config
    except FileNotFoundError:
        raise FileNotFoundError(f"配置文件 {config_path} 不存在")
    except json.JSONDecodeError:
        raise ValueError("配置文件格式错误，必须是合法的 JSON")

def main():
    # 解析命令行参数
    parser = argparse.ArgumentParser(description='生成漏洞报告文档')
    parser.add_argument('md_file_path', help='Markdown 文件路径（如：漏洞描述.md）')
    args = parser.parse_args()

    try:
        # 定义报告目录（使用绝对路径更可靠）
        report_dir = os.path.abspath('report')  # 关键修复点1：转换为绝对路径
        os.makedirs(report_dir, exist_ok=True)  # 确保目录存在
        # 读取配置文件
        config = read_config()
        xls_file_path = config['xls_file_path']
        word_template_path = config['word_template_path']

        # 验证文件存在性
        if not os.path.exists(args.md_file_path):
            raise FileNotFoundError(f"Markdown 文件 {args.md_file_path} 不存在")
        if not os.path.exists(xls_file_path):
            raise FileNotFoundError(f"Excel 文件 {xls_file_path} 不存在")
        if not os.path.exists(word_template_path):
            raise FileNotFoundError(f"Word 模板文件 {word_template_path} 不存在")

        # 原有业务逻辑
        md_content = read_md_file(args.md_file_path)
        md_result = parse_md_string(md_content)
        find_vulnerability_in_excel(md_result, xls_file_path)
        word_replacements = flatten_dict(md_result)
        word_output_file = word_replacements['{公司名称}'] + word_replacements['{系统名称}'] + "存在" + word_replacements['{漏洞名称}'] + ".docx"
        word_output_file = os.path.join(report_dir, word_output_file)
        replace_in_word_template(word_template_path, word_output_file, word_replacements)

        print(f"报告生成成功：{word_output_file}")

    except Exception as e:
        print(f"错误发生: {str(e)}")
        exit(1)

if __name__ == "__main__":
    main()

